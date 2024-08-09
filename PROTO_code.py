import psutil
import win32gui
import time
import csv
import pandas as pd
import tkinter as tk
from tkinter import messagebox, Toplevel, Label, OptionMenu, StringVar, Button
from threading import Thread, Event
import os
from sklearn.preprocessing import StandardScaler
from sklearn.ensemble import IsolationForest
from datetime import datetime
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np

# Excel file setup
EXCEL_PATH = r'E:\2024\Machine Learning\Personal progress Optimzer\All code\proto\Master_excel_sheet.xlsx'
SHEET_NAME = 'Sheet1'  # Adjust if your sheet has a different name

def load_master_database():
    """Load the master database from Excel into a DataFrame."""
    return pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME)

def save_to_master_database(app_name, category):
    """Save a new application and its category to the Excel master database."""
    df = load_master_database()
    new_entry = pd.DataFrame({'Application': [app_name], 'Category': [category]})
    df = pd.concat([df, new_entry], ignore_index=True)
    df.to_excel(EXCEL_PATH, sheet_name=SHEET_NAME, index=False)

# Global variables
monitoring_event = Event()
LOG_FILE = None
ANALYSIS_FOLDER = None

# Define window title categories
CATEGORIES = {
    'gaming': ['game', 'games', 'play', 'arcade', 'fifa'],
    'movies': ['netflix', 'prime', 'hulu', 'video', 'movie'],
    'music': ['spotify', 'music', 'soundcloud'],
    'social_media': ['facebook', 'instagram', 'twitter', 'linkedin'],
    'productivity': ['excel', 'word', 'powerpoint', 'editor'],
    'design': ['photoshop', 'autocad', 'illustrator']
}

COMMON_CATEGORIES = [
    'gaming', 'movies', 'music', 'social_media',
    'productivity', 'design', 'other'
]

def categorize_application(app_name):
    """Prompt the user to categorize an application with a dropdown list."""
    def on_select_category():
        selected_category = category_var.get()
        save_to_master_database(app_name, selected_category)
        prompt_window.destroy()

    prompt_window = Toplevel(root)
    prompt_window.title("Categorize Application")
    prompt_window.geometry("300x150")
    prompt_window.grab_set()

    Label(prompt_window, text=f"Categorize the application '{app_name}'").pack(pady=10)

    category_var = StringVar(prompt_window)
    category_var.set(COMMON_CATEGORIES[-1])  # Default to 'other'
    dropdown = OptionMenu(prompt_window, category_var, *COMMON_CATEGORIES)
    dropdown.pack(pady=10)

    Button(prompt_window, text="Save", command=on_select_category).pack(pady=10)

    root.wait_window(prompt_window)
    return category_var.get()

def get_application_category(app_name):
    """Get the category of an application from the master database or prompt the user."""
    df = load_master_database()
    if app_name in df['Application'].values:
        return df.loc[df['Application'] == app_name, 'Category'].values[0]
    else:
        return categorize_application(app_name)

def create_analysis_folder():
    """Create a folder named 'analysis_datetime' in the code's directory."""
    global ANALYSIS_FOLDER
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    folder_name = f"analysis_{timestamp}"
    ANALYSIS_FOLDER = os.path.join(os.path.dirname(__file__), folder_name)
    os.makedirs(ANALYSIS_FOLDER, exist_ok=True)
    global LOG_FILE
    LOG_FILE = os.path.join(ANALYSIS_FOLDER, 'activity_log.csv')
    update_status(f"Log file will be saved to: {LOG_FILE}")

def log_data(active_window, cpu_usage):
    """Log the data to the CSV file with headers if the file is new."""
    if LOG_FILE:
        file_exists = os.path.isfile(LOG_FILE)
        with open(LOG_FILE, mode='a', newline='') as file:
            writer = csv.writer(file)
            if not file_exists:
                writer.writerow(['Timestamp', 'Active Window', 'CPU Usage'])
            writer.writerow([time.strftime("%Y-%m-%d %H:%M:%S"), active_window, cpu_usage])

def get_active_window_title():
    """Get the title of the currently active window."""
    try:
        hwnd = win32gui.GetForegroundWindow()
        title = win32gui.GetWindowText(hwnd)
        return title if title else "No active window"
    except Exception as e:
        return f"Error: {e}"

def classify_title(title):
    """Classify window title using the master database or ask the user for a category."""
    title_lower = title.lower()
    category = get_application_category(title_lower)
    if category:
        return category
    else:
        for category, keywords in CATEGORIES.items():
            if any(keyword in title_lower for keyword in keywords):
                return category
        return "other"

def classify_window_titles(df):
    """Classify window titles to identify specific categories."""
    df['Classification'] = df['Active Window'].apply(classify_title)
    return df

def monitor_activity():
    """Monitor and log system activity."""
    while monitoring_event.is_set():
        active_window_title = get_active_window_title()
        cpu_usage = psutil.cpu_percent(interval=1)
        log_data(active_window_title, cpu_usage)
        time.sleep(10)

def start_monitoring():
    """Start monitoring activity."""
    create_analysis_folder()
    monitoring_event.set()
    Thread(target=monitor_activity, daemon=True).start()
    update_status("Monitoring started.")

def stop_monitoring():
    """Stop monitoring activity."""
    monitoring_event.clear()
    analyze_data()
    update_status("Monitoring stopped.")

def generate_ai_advice(row):
    """Generate advice based on application usage."""
    advice = ""
    if row['Anomaly'] == -1:
        advice += "Anomalous usage detected. "
    
    if row['Classification'] == 'gaming':
        advice += "Consider taking regular breaks during gaming sessions. "
    elif row['Classification'] == 'movies':
        advice += "Limit movie watching to maintain productivity. "
    elif row['Classification'] == 'music':
        advice += "Enjoy music in moderation to avoid distractions. "
    elif row['Classification'] == 'social_media':
        advice += "Avoid excessive time on social media. "
    elif row['Classification'] == 'productivity':
        advice += "Keep up the good work! Stay focused and maintain productivity. "
    elif row['Classification'] == 'design':
        advice += "Great job on creative tasks! Remember to balance with other activities. "
    elif row['Classification'] == 'other':
        advice += "Consider categorizing unrecognized applications or activities."

    return advice.strip()

def load_master_database():
    """Load the master database from Excel into a DataFrame, creating columns if they don't exist."""
    try:
        df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME)
        print("Columns in master database:", df.columns)
        
        # Ensure 'Application' and 'Category' columns exist
        if 'Application' not in df.columns:
            df['Application'] = np.nan
        if 'Category' not in df.columns:
            df['Category'] = np.nan
        
        return df
    except Exception as e:
        print(f"Error loading master database: {e}")
        # Return a DataFrame with required columns if loading fails
        return pd.DataFrame(columns=['Application', 'Category'])

def save_to_master_database(app_name, category):
    """Save a new application and its category to the Excel master database."""
    df = load_master_database()
    
    # Append new data
    new_entry = pd.DataFrame([[app_name, category]], columns=['Application', 'Category'])
    df = pd.concat([df, new_entry], ignore_index=True)
    
    # Save the updated DataFrame back to Excel
    try:
        with pd.ExcelWriter(EXCEL_PATH, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
            df.to_excel(writer, sheet_name=SHEET_NAME, index=False)
    except Exception as e:
        print(f"Error saving to master database: {e}")

# Ensure 'Application' and 'Category' columns are created if missing
df_master = load_master_database()
print(df_master.head())

def analyze_data():
    """Analyze the logged data to provide insights and recommendations."""
    if not LOG_FILE:
        messagebox.showwarning("Warning", "Log file not found. Please ensure monitoring is started.")
        return

    try:
        df = pd.read_csv(LOG_FILE)
        df['Timestamp'] = pd.to_datetime(df['Timestamp'])
        df['Date'] = df['Timestamp'].dt.date
        df['Time'] = df['Timestamp'].dt.time
        df['Active Window'] = df['Active Window'].str.strip()

        # Print DataFrame for debugging
        print("Columns in activity log:", df.columns)
        
        # Calculate duration for each window
        df['Duration'] = df.groupby('Active Window')['Timestamp'].transform(
            lambda x: (x.max() - x.min()).total_seconds() / 60
        )
        df = df.drop_duplicates(subset=['Active Window'])
        df = df[['Active Window', 'Duration']]
        df = df.groupby('Active Window').sum().reset_index()

        # Check if 'Application' column is present
        df_master = load_master_database()
        if 'Application' not in df_master.columns or 'Category' not in df_master.columns:
            raise ValueError("Excel file must contain 'Application' and 'Category' columns.")

        # Classify the active windows
        df = classify_window_titles(df)

        # AI Integration: Detecting anomalies
        scaler = StandardScaler()
        df['Normalized Duration'] = scaler.fit_transform(df[['Duration']])
        model = IsolationForest(contamination=0.1)
        df['Anomaly'] = model.fit_predict(df[['Normalized Duration']])
        
        # Generate advice for each application
        df['AI_Advice'] = df.apply(generate_ai_advice, axis=1)

        # Assess computer goodness
        avg_cpu_usage = psutil.cpu_percent(interval=1)
        total_anomalies = df['Anomaly'].sum()

        # Get CPU temperature if available
        try:
            temp = psutil.sensors_temperatures()
            if 'coretemp' in temp:
                cpu_temp = temp['coretemp'][0].current
            else:
                cpu_temp = "Not available"
        except AttributeError:
            cpu_temp = "Not available"

        computer_goodness = "Good" if total_anomalies == 0 and avg_cpu_usage < 50 else "Needs Attention"

        # Assess human goodness
        productivity_apps = ['excel', 'word', 'powerpoint']
        entertainment_apps = ['youtube', 'netflix']
        design_apps = ['photoshop', 'autocad', 'illustrator']
        gaming_apps = df[df['Classification'] == 'gaming']['Duration'].sum()
        productive_time = df[df['Active Window'].str.contains('|'.join(productivity_apps), case=False, na=False)]['Duration'].sum()
        entertainment_time = df[df['Active Window'].str.contains('|'.join(entertainment_apps), case=False, na=False)]['Duration'].sum()
        design_time = df[df['Active Window'].str.contains('|'.join(design_apps), case=False, na=False)]['Duration'].sum()
        total_time = df['Duration'].sum()
        human_goodness = "Balanced" if productive_time >= entertainment_time else "Needs Improvement"

        # Prepare the analysis report in a Word document
        doc = Document()
        doc.add_heading('Activity Analysis Report', 0)

        doc.add_heading('Computer Goodness:', level=1)
        doc.add_paragraph(f"Computer Goodness: {computer_goodness}")
        doc.add_paragraph(f"Average CPU Usage: {avg_cpu_usage}%")
        doc.add_paragraph(f"CPU Temperature: {cpu_temp}")
        doc.add_paragraph(f"Total Anomalies Detected: {total_anomalies}")

        doc.add_heading('Human Goodness:', level=1)
        doc.add_paragraph(f"Human Goodness: {human_goodness}")
        doc.add_paragraph(f"Total Productive Time: {productive_time:.2f} minutes")
        doc.add_paragraph(f"Total Entertainment Time: {entertainment_time:.2f} minutes")
        doc.add_paragraph(f"Total Design Software Time: {design_time:.2f} minutes")
        doc.add_paragraph(f"Total Gaming Time: {gaming_apps:.2f} minutes")
        doc.add_paragraph(f"Total Time Spent on Computer: {total_time:.2f} minutes")

        doc.add_heading('Detailed Analysis:', level=1)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Application'
        hdr_cells[1].text = 'Time Spent (minutes)'
        hdr_cells[2].text = 'Anomaly Detected'
        hdr_cells[3].text = 'Classification'
        hdr_cells[4].text = 'AI Advice'

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row['Active Window']
            row_cells[1].text = f"{row['Duration']:.2f}"
            row_cells[2].text = 'Yes' if row['Anomaly'] == -1 else 'No'
            row_cells[3].text = row['Classification']
            row_cells[4].text = row['AI_Advice']

        # Add a pie chart with improved clarity
        doc.add_heading('Usage Pie Chart:', level=1)
        pie_chart_path = os.path.join(ANALYSIS_FOLDER, 'usage_pie_chart.png')
        plt.figure(figsize=(10, 8))

        category_colors = {
            'gaming': 'lightblue',
            'movies': 'lightgreen',
            'music': 'lightcoral',
            'social_media': 'lightsalmon',
            'productivity': 'lightpink',
            'design': 'lightyellow',
            'other': 'lightgray'
        }
        
        category_summary = df.groupby('Classification')['Duration'].sum()
        plt.pie(category_summary, labels=category_summary.index, colors=[category_colors.get(cat, 'lightgray') for cat in category_summary.index],
                autopct='%1.1f%%', startangle=140)
        plt.title('Time Spent on Different Activities')
        plt.savefig(pie_chart_path)
        plt.close()
        doc.add_picture(pie_chart_path, width=Inches(5))

        doc_path = os.path.join(ANALYSIS_FOLDER, 'analysis_report.docx')
        doc.save(doc_path)
        update_status(f"Analysis report saved to: {doc_path}")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred during data analysis: {e}")

def update_status(message):
    """Update the status label in the GUI."""
    status_label.config(text=message)

# GUI Setup
root = tk.Tk()
root.title("Activity Monitor")
root.geometry("400x200")

# Start and Stop buttons
start_button = tk.Button(root, text="Start Monitoring", command=start_monitoring)
start_button.pack(pady=10)

stop_button = tk.Button(root, text="Stop Monitoring", command=stop_monitoring)
stop_button.pack(pady=10)

# Status label
status_label = tk.Label(root, text="Monitoring status: Not started")
status_label.pack(pady=10)

# Run the Tkinter event loop
root.mainloop()
